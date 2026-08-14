import pandas as pd
import matplotlib.pyplot as plt
import matplotlib as mpl
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

mpl.rcParams["font.family"] = "Malgun Gothic"
mpl.rcParams["axes.unicode_minus"] = False

df = pd.read_csv("data/핀테크_정제완료.csv", encoding="utf-8-sig")
grouped = df.groupby("가맹점업종")["거래금액"].sum().sort_values(ascending=False)
total = df["거래금액"].sum()
top3_share = grouped.head(3).sum() / total * 100

# 차트 저장
fig, ax = plt.subplots(figsize=(9, 5))
bars = ax.bar(grouped.index, grouped.values, color="#4C72B0")
ax.set_title("가맹점업종별 총 거래금액")
ax.set_xlabel("가맹점업종")
ax.set_ylabel("총 거래금액 (원)")
ax.tick_params(axis="x", rotation=30)
for bar in bars:
    height = bar.get_height()
    ax.annotate(f"{height:,.0f}", xy=(bar.get_x() + bar.get_width() / 2, height),
                xytext=(0, 3), textcoords="offset points", ha="center", va="bottom", fontsize=8)
ax.margins(y=0.12)
fig.tight_layout()
chart_path = "scripts/_chart_가맹점업종.png"
fig.savefig(chart_path, dpi=150)
plt.close(fig)

# docx 작성
doc = Document()

title = doc.add_heading("가맹점업종별 거래금액 분석 보고서", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("데이터 출처: data/핀테크_정제완료.csv (정제 완료본, 11,738건)").italic = True

doc.add_picture(chart_path, width=Inches(6))
last_p = doc.paragraphs[-1]
last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("분석 요약", level=2)

top = grouped.index[0]
bottom = grouped.index[-1]

bullets = [
    f"전체 거래금액 합계는 {total:,.0f}원이며, 9개 가맹점업종 중 '{top}' 업종이 {grouped.iloc[0]:,.0f}원(전체의 {grouped.iloc[0]/total*100:.1f}%)으로 가장 큰 비중을 차지합니다.",
    f"'쇼핑', '교육' 업종이 뒤를 이어 상위 3개 업종(여행·쇼핑·교육)이 전체 거래금액의 {top3_share:.1f}%를 차지 — 소수 업종에 매출이 집중되어 있습니다.",
    f"거래 건수 기준으로는 '식음료'(2,630건), '쇼핑'(2,300건)이 가장 빈번하지만, 건당 평균 금액이 낮아(각 18,337원, 52,696원) 총액 순위는 상대적으로 낮습니다. 반면 '여행'은 건수(559건)는 적어도 건당 평균이 229,512원으로 압도적으로 높아 총액 1위를 견인했습니다.",
    f"'교통' 업종은 총 거래금액이 {grouped.iloc[-1]:,.0f}원(전체의 {grouped.iloc[-1]/total*100:.1f}%)으로 가장 낮으며, 건당 평균 금액도 3,542원으로 소액·고빈도 결제 패턴을 보입니다.",
    "종합하면 매출 규모는 '거래 빈도'보다 '건당 단가'에 더 크게 좌우되는 구조이며, 여행·교육처럼 건당 단가가 높은 업종에 대한 프로모션/제휴가 매출 확대에 효과적일 수 있습니다.",
]

for b in bullets:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(b)

doc.save("가맹점업종_거래금액_보고서.docx")
print("저장 완료: 가맹점업종_거래금액_보고서.docx")
